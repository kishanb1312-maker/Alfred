#!/usr/bin/env python3
"""Alfred · alfred — the host-agnostic setup/maintenance CLI (§5).

A single entrypoint for the plugin's first-run setup and health checks. It is the
ONLY place secrets ever enter Alfred, and they enter via getpass (no echo) —
never typed into chat, never printed, never written by the assistant.

Subcommands:
    alfred bootstrap                 create the venv + install deps + init (cross-platform, no shell)
    alfred env                       print the venv's python path for THIS OS (copy-pasteable)
    alfred init                      create ~/.alfred/{config,data,output}/ + copy templates
    alfred set-secret <KEY>          getpass (hidden) → write KEY=value into ~/.alfred/.env
    alfred detect-telegram-chat-id   getUpdates → write TELEGRAM_CHAT_ID into .env
    alfred doctor                    full health checklist (superset of orchestrate --preflight)
    alfred parse-resume <path.docx>  copy resume into data/ + print extracted vs still-needed fields

Cross-platform note: this CLI runs the same on Windows, macOS, and Linux — every
path is built with os.path and the home dir resolves via os.path.expanduser("~")
(=%USERPROFILE% on Windows). `bootstrap` exists so first-run setup needs no Bash:
it drives `python -m venv` + pip via subprocess in pure Python.

Path model (see paths.py §4): CODE lives in the plugin dir ($CLAUDE_PLUGIN_ROOT),
USER DATA lives in the home dir. Every write here targets `home_target()` — the real
$ALFRED_HOME or ~/.alfred — never the legacy repo-root fallback, because these
commands create/manage that home dir (on first `init` it does not exist yet).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
from typing import Any, Callable, Dict, List, Optional

# scripts/ on sys.path (sys.path[0] when run directly; already present when imported).
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import paths  # noqa: E402  single source of truth for CODE_ROOT + example templates


# ---------------------------------------------------------------------------
# Key sets (from .env.example / SPEC §8a)
# ---------------------------------------------------------------------------

REQUIRED_SECRETS = ["TELEGRAM_BOT_TOKEN", "TELEGRAM_CHAT_ID",
                    "GMAIL_ADDRESS", "GMAIL_APP_PASSWORD"]
OPTIONAL_SECRETS = ["HUNTER_API_KEY", "APOLLO_API_KEY",
                    "NOTION_API_KEY", "NOTION_DATABASE_ID"]
KNOWN_KEYS = REQUIRED_SECRETS + OPTIONAL_SECRETS


# ---------------------------------------------------------------------------
# Home-dir resolution — the CLI's write target (NOT paths.DATA_HOME; see docstring)
# ---------------------------------------------------------------------------

def home_target() -> str:
    """Resolved the same way as paths._resolve_data_home() (minus the repo-root
    fallback) so reads and writes always agree — including for a pre-rename install
    that still lives in $WARMAPPLY_HOME / ~/.warmapply."""
    for var in ("ALFRED_HOME", "WARMAPPLY_HOME"):
        override = os.environ.get(var, "").strip()
        if override:
            return os.path.abspath(os.path.expanduser(override))

    fresh = os.path.join(os.path.expanduser("~"), ".alfred")
    legacy = os.path.join(os.path.expanduser("~"), ".warmapply")
    if not os.path.isdir(fresh) and os.path.isdir(legacy):
        return legacy
    return fresh


def home_paths(home: Optional[str] = None) -> Dict[str, str]:
    home = home or home_target()
    return {
        "home": home,
        "config": os.path.join(home, "config"),
        "data": os.path.join(home, "data"),
        "output": os.path.join(home, "output"),
        "env": os.path.join(home, ".env"),
        "search": os.path.join(home, "config", "search.yaml"),
        "profile": os.path.join(home, "config", "profile.yaml"),
    }


# ---------------------------------------------------------------------------
# .env upsert — PURE text transform (unit-testable, never touches getpass/stdin)
# ---------------------------------------------------------------------------

def upsert_env_line(text: str, key: str, value: str) -> str:
    """Return `text` with `KEY=value` set: replace the first existing KEY line
    (commented or not, with/without `export`) else append. Order & comments kept."""
    pat = re.compile(rf"^\s*(?:export\s+)?#?\s*{re.escape(key)}\s*=")
    out: List[str] = []
    replaced = False
    for line in text.splitlines():
        if not replaced and pat.match(line):
            out.append(f"{key}={value}")
            replaced = True
        else:
            out.append(line)
    if not replaced:
        out.append(f"{key}={value}")
    return "\n".join(out) + "\n"


def write_secret(env_path: str, key: str, value: str) -> None:
    """Atomically upsert KEY into the env file with 0600 perms. Value never logged."""
    os.makedirs(os.path.dirname(env_path), exist_ok=True)
    text = ""
    if os.path.isfile(env_path):
        with open(env_path, "r", encoding="utf-8") as fh:
            text = fh.read()
    new_text = upsert_env_line(text, key, value)
    tmp = env_path + ".tmp"
    # Create the temp file 0600 from the START (not chmod-after-write) so a secret
    # is never briefly world-readable. O_CREAT applies the mode only on creation, so
    # drop any stale tmp first to guarantee our perms win.
    if os.path.exists(tmp):
        os.unlink(tmp)
    fd = os.open(tmp, os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    with os.fdopen(fd, "w", encoding="utf-8") as fh:
        fh.write(new_text)
    os.replace(tmp, env_path)


def _read_env_values(env_path: str) -> Dict[str, str]:
    """Parse KEY=value pairs from the env file (for doctor's filled/blank check)."""
    values: Dict[str, str] = {}
    if not os.path.isfile(env_path):
        return values
    with open(env_path, "r", encoding="utf-8") as fh:
        for raw in fh:
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            if line.startswith("export "):
                line = line[len("export "):].lstrip()
            k, _, v = line.partition("=")
            values[k.strip()] = v.strip().strip("'\"")
    return values


# ---------------------------------------------------------------------------
# venv location — identical logic on every OS, only the leaf dir differs
# ---------------------------------------------------------------------------

def venv_dir(home: Optional[str] = None) -> str:
    return os.path.join(home or home_target(), ".venv")


def venv_python(home: Optional[str] = None) -> str:
    """Path to the venv's python interpreter for the CURRENT OS.

    Windows puts it at .venv\\Scripts\\python.exe; POSIX at .venv/bin/python.
    """
    vdir = venv_dir(home)
    if os.name == "nt":
        return os.path.join(vdir, "Scripts", "python.exe")
    return os.path.join(vdir, "bin", "python")


# ---------------------------------------------------------------------------
# env — print the venv python path (so the wizard can hand over one command)
# ---------------------------------------------------------------------------

def cmd_env(args: argparse.Namespace) -> int:
    vpy = venv_python()
    exists = os.path.isfile(vpy)
    print(vpy if exists else f"{vpy}  (not created yet — run: alfred bootstrap)")
    return 0 if exists else 1


# ---------------------------------------------------------------------------
# bootstrap — venv + deps + init, in pure Python (no Bash / PowerShell needed)
# ---------------------------------------------------------------------------

def cmd_bootstrap(args: argparse.Namespace) -> int:
    """Create the venv, install requirements into it, and run init.

    Runs the same on Windows, macOS, and Linux — it shells out only to Python
    itself (venv + pip) via subprocess, so no OS-specific shell syntax is needed.
    """
    import subprocess

    hp = home_paths()
    os.makedirs(hp["home"], exist_ok=True)
    vdir = venv_dir(hp["home"])
    vpy = venv_python(hp["home"])

    # 1) venv (idempotent — python -m venv is a no-op refresh if it already exists)
    if os.path.isfile(vpy):
        print(f"  •  venv exists : {vdir}")
    else:
        print(f"  ⏳ creating venv: {vdir}")
        rc = subprocess.call([sys.executable, "-m", "venv", vdir])
        if rc != 0 or not os.path.isfile(vpy):
            print("❌ failed to create the virtual environment. Is Python 3.8+ installed "
                  "and on PATH? (Windows: install from python.org and tick "
                  "'Add python.exe to PATH'.)")
            return 1

    # 2) deps — install requirements.txt into the venv (not the system interpreter)
    req = os.path.join(paths.CODE_ROOT, "requirements.txt")
    if not os.path.isfile(req):
        print(f"⚠️  requirements.txt not found at {req} — skipping dependency install.")
    else:
        print(f"  ⏳ installing deps from {req}")
        rc = subprocess.call([vpy, "-m", "pip", "install", "-q",
                              "--disable-pip-version-check", "-r", req])
        if rc != 0:
            print("❌ dependency install failed. Re-run manually:\n"
                  f'   "{vpy}" -m pip install -r "{req}"')
            return 1

    # 3) init the data dir + templates
    print("  ⏳ initializing data dir")
    cmd_init(args)

    print("-" * 60)
    print("✅ bootstrap complete. From here, run the CLI with the VENV python:")
    print(f'   "{vpy}" "{os.path.join(paths.CODE_ROOT, "scripts", "alfred_cli.py")}" doctor')
    return 0


# ---------------------------------------------------------------------------
# init
# ---------------------------------------------------------------------------

def cmd_init(args: argparse.Namespace) -> int:
    hp = home_paths()
    created: List[str] = []
    skipped: List[str] = []

    for d in (hp["home"], hp["config"], hp["data"], hp["output"]):
        if os.path.isdir(d):
            skipped.append(d)
        else:
            os.makedirs(d, exist_ok=True)
            created.append(d)

    # Copy bundled templates (never overwrite a file the user already has → idempotent).
    for src, dst in ((paths.search_example(), hp["search"]),
                     (paths.profile_example(), hp["profile"]),
                     (paths.env_example(), hp["env"])):
        if os.path.exists(dst):
            skipped.append(dst)
            continue
        if not os.path.exists(src):
            print(f"  ⚠️  template not found (skipped): {src}")
            continue
        shutil.copyfile(src, dst)
        if dst == hp["env"]:
            os.chmod(dst, 0o600)  # .env holds secrets → lock it down from creation
        created.append(dst)

    print(f"Alfred · init  →  {hp['home']}")
    print("-" * 60)
    for c in created:
        print(f"  ✅ created : {c}")
    for s in skipped:
        print(f"  •  exists  : {s}")
    print("-" * 60)
    print("Next:")
    print("  1. alfred set-secret TELEGRAM_BOT_TOKEN   (then GMAIL_ADDRESS, GMAIL_APP_PASSWORD)")
    print("  2. message your bot once, then: alfred detect-telegram-chat-id")
    print("  3. edit config/search.yaml + config/profile.yaml (or use /alfred-setup)")
    print("  4. alfred doctor")
    return 0


# ---------------------------------------------------------------------------
# set-secret  (the ONLY way secrets enter — getpass, no echo)
# ---------------------------------------------------------------------------

def cmd_set_secret(args: argparse.Namespace) -> int:
    import getpass

    key = args.key.strip()
    if not re.fullmatch(r"[A-Z][A-Z0-9_]*", key):
        print(f"❌ invalid key '{key}' — expected UPPER_SNAKE_CASE (e.g. TELEGRAM_BOT_TOKEN)")
        return 2
    if key not in KNOWN_KEYS:
        print(f"⚠️  '{key}' is not a recognized Alfred key — setting it anyway.")

    hp = home_paths()
    if not os.path.isdir(hp["home"]):
        print(f"⚠️  {hp['home']} does not exist yet — creating it "
              "(run `alfred init` for the full skeleton).")
        os.makedirs(hp["home"], exist_ok=True)

    try:
        value = getpass.getpass(f"Enter value for {key} (input hidden, not echoed): ")
    except (EOFError, KeyboardInterrupt):
        print("\n(aborted — nothing written)")
        return 1

    if value == "":
        print("(empty value — nothing written)")
        return 1

    write_secret(hp["env"], key, value)
    del value  # do not keep the secret in memory longer than needed
    print(f"✅ {key} saved to {hp['env']}  (value hidden; file perms 600).")
    return 0


# ---------------------------------------------------------------------------
# detect-telegram-chat-id
# ---------------------------------------------------------------------------

def extract_chat_id(updates_json: Any) -> Optional[str]:
    """PURE: newest chat.id from a Telegram getUpdates response, or None."""
    if not isinstance(updates_json, dict) or not updates_json.get("ok"):
        return None
    result = updates_json.get("result") or []
    for update in reversed(result):
        if not isinstance(update, dict):
            continue
        for key in ("message", "edited_message", "channel_post", "my_chat_member"):
            msg = update.get(key)
            if isinstance(msg, dict):
                chat = msg.get("chat") or {}
                cid = chat.get("id")
                if cid is not None:
                    return str(cid)
    return None


def _get_updates(token: str, timeout: int = 10) -> Any:
    """Network chokepoint: fetch getUpdates. Injected/faked in the dry-run."""
    import requests
    url = f"https://api.telegram.org/bot{token}/getUpdates"
    resp = requests.get(url, params={"timeout": 0}, timeout=timeout)
    resp.raise_for_status()
    return resp.json()


def _redact(text: str, secret: str) -> str:
    """Strip a secret from a string before it is printed/logged. The bot token is
    embedded in the request URL, so requests' HTTPError/ConnectionError/Timeout
    messages all echo it — redact it from EVERY exception, not just some."""
    if secret and text:
        text = text.replace(secret, "***REDACTED***")
    return text


def cmd_detect(args: argparse.Namespace, _fetch: Optional[Callable[[str], Any]] = None) -> int:
    hp = home_paths()
    paths.load_env(hp["env"])
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "").strip()
    if not token:
        print("❌ TELEGRAM_BOT_TOKEN not set. First run:  alfred set-secret TELEGRAM_BOT_TOKEN")
        return 1

    fetch = _fetch or _get_updates
    try:
        data = fetch(token)
    except Exception as exc:  # network / HTTP error — its message may embed the token
        print(f"❌ could not reach Telegram getUpdates: {_redact(str(exc), token)}")
        return 1

    cid = extract_chat_id(data)
    if not cid:
        print("No message found yet. In Telegram, send ANY message to your bot, then re-run:")
        print("    alfred detect-telegram-chat-id")
        return 1

    write_secret(hp["env"], "TELEGRAM_CHAT_ID", cid)  # chat id is an identifier, not a secret
    print(f"✅ TELEGRAM_CHAT_ID = {cid}  saved to {hp['env']}.")
    return 0


# ---------------------------------------------------------------------------
# doctor  (superset of orchestrate --preflight)
# ---------------------------------------------------------------------------

_DEP_MODULES = [("yaml", "PyYAML"), ("docx", "python-docx"), ("requests", "requests"),
                ("dns.resolver", "dnspython"), ("email_validator", "email-validator")]


def _mark(status: str) -> str:
    return {"ok": "✅", "warn": "⚠️ ", "fail": "❌"}.get(status, "•")


def cmd_doctor(args: argparse.Namespace) -> int:
    import orchestrate  # reuse the exact preflight checks (configs/resume/dry_run/pause)

    hp = home_paths()
    if not os.path.isdir(hp["home"]):
        print(f"❌ Alfred is not initialized — run `alfred init` "
              f"(expected {hp['home']}).")
        return 1
    paths.load_env(hp["env"])

    checks: List[tuple] = []  # (status, label, detail)

    # 1) Python deps importable (required)
    missing = [pkg for mod, pkg in _DEP_MODULES if not _importable(mod)]
    checks.append(("fail" if missing else "ok", "Python deps",
                   (", ".join(missing) + " missing — pip install -r requirements.txt")
                   if missing else "all importable"))

    # 2) Configs valid + resume present + dry_run/pause — delegate to preflight
    report = orchestrate.preflight(
        search_path=hp["search"], profile_path=hp["profile"],
        resume_glob=os.path.join(hp["data"], "*.docx"),
        pause_flag=os.path.join(hp["data"], "paused.flag"),
    )
    c = report["checks"]
    checks.append(("ok" if c["search_config"]["exists"] else "fail",
                   "search config", hp["search"]))
    checks.append(("ok" if c["profile_config"]["exists"] else "fail",
                   "profile config", hp["profile"]))
    checks.append(("ok" if c["resume"]["exists"] else "fail", "master resume",
                   (c["resume"]["found"][0] if c["resume"]["found"] else c["resume"]["glob"])))

    # 3) Required secrets filled (required)
    env_vals = _read_env_values(hp["env"])
    empty_required = [k for k in REQUIRED_SECRETS if not env_vals.get(k, "").strip()]
    checks.append(("fail" if empty_required else "ok", "required secrets",
                   ("blank: " + ", ".join(empty_required) + "  (alfred set-secret <KEY>)")
                   if empty_required else "all set (TELEGRAM_*, GMAIL_*)"))

    # 4) Optional secrets (advisory)
    empty_optional = [k for k in OPTIONAL_SECRETS if not env_vals.get(k, "").strip()]
    checks.append(("warn" if empty_optional else "ok", "optional secrets",
                   ("blank (fine to leave): " + ", ".join(empty_optional))
                   if empty_optional else "all set"))

    # 5) LibreOffice / soffice (advisory — needed only for DOCX→PDF at tailoring).
    #    Use the cross-platform finder so a Windows install off PATH is still detected.
    #    Finding the binary is NOT enough: an install missing the Writer module still
    #    ships `soffice` but cannot open a .docx, so probe Writer too or this check
    #    reports green while every resume/cover-letter PDF silently fails to build.
    writer: Optional[bool] = None
    try:
        import docx_to_pdf
        soffice = docx_to_pdf.find_soffice()
        writer = docx_to_pdf.writer_available()
    except Exception:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        checks.append(("warn", "LibreOffice (soffice)",
                       "not found — install for PDF export: "
                       "winget install TheDocumentFoundation.LibreOffice  /  "
                       "brew install --cask libreoffice  /  apt-get install libreoffice"))
    elif writer is False:
        checks.append(("warn", "LibreOffice (soffice)",
                       f"{soffice} — but the Writer module is MISSING, so .docx→PDF will "
                       "fail and no resume/cover letter will be attached. Install it: "
                       "apt-get install libreoffice-writer  (or the full libreoffice package)"))
    else:
        checks.append(("ok", "LibreOffice (soffice)",
                       soffice + ("" if writer else "  (Writer module unverified)")))

    # 6) Notion MCP (advisory — reached via Claude Code connector, not .env)
    if env_vals.get("NOTION_API_KEY", "").strip():
        checks.append(("ok", "Notion", "direct NOTION_API_KEY present"))
    else:
        checks.append(("warn", "Notion",
                       "via Claude Code Notion MCP connector — verify with /mcp in Claude Code"))

    # ---- render ----
    print("Alfred · doctor")
    print("-" * 60)
    for status, label, detail in checks:
        print(f"  {_mark(status)} {label:<18}: {detail}")

    dry = report["dry_run"]
    dry_show = "ON (no real sends/submits)" if dry else ("OFF (LIVE)" if dry is False else "unknown")
    print(f"  •  {'dry_run':<18}: {dry_show}")
    print(f"  {'⛔' if report['paused'] else '▶️ '} {'/pause':<18}: "
          f"{'SET — application-agent halted' if report['paused'] else 'not set'}")
    print("-" * 60)

    has_fail = any(s == "fail" for s, _, _ in checks) or report["paused"]
    if has_fail:
        print("RESULT: NOT READY — resolve the ❌ items above (⚠️ items are advisory).")
        return 1
    warns = [l for s, l, _ in checks if s == "warn"]
    if warns:
        print(f"RESULT: READY ✅  (advisories: {', '.join(warns)})")
    else:
        print("RESULT: READY ✅  — all green.")
    if dry:
        print("  (dry_run is ON — a full pass prepares everything but sends/submits nothing.)")
    return 0


def _importable(mod: str) -> bool:
    try:
        __import__(mod)
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# parse-resume
# ---------------------------------------------------------------------------

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)")
_URL_RE = re.compile(r"https?://[^\s)>\]]+", re.IGNORECASE)

_SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "go", "golang", "rust", "ruby",
    "c++", "c#", "php", "scala", "kotlin", "swift", "bash", "shell",
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible", "linux",
    "jenkins", "gitlab", "github actions", "ci/cd", "git",
    "sql", "postgresql", "postgres", "mysql", "mongodb", "redis", "kafka",
    "elasticsearch", "spark", "hadoop", "airflow",
    "react", "angular", "vue", "node", "node.js", "django", "flask", "fastapi",
    "spring", "rails", ".net", "graphql", "rest", "grpc",
    "prometheus", "grafana", "datadog", "sre", "devops",
]
_TITLE_KEYWORDS = [
    "software engineer", "senior software engineer", "backend engineer",
    "frontend engineer", "full stack engineer", "full-stack engineer",
    "site reliability engineer", "sre", "devops engineer", "platform engineer",
    "data engineer", "cloud engineer", "systems engineer", "systems administrator",
    "solutions architect", "software architect", "developer", "programmer",
    "engineering manager", "tech lead", "team lead",
]


def read_docx_text(path: str) -> str:
    import docx
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:  # skills often live in tables
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_from_resume(text: str) -> Dict[str, Any]:
    """Best-effort structured pull from raw resume text (no invention)."""
    emails = sorted(set(_EMAIL_RE.findall(text)))
    phones = sorted({re.sub(r"\s+", " ", p).strip() for p in _PHONE_RE.findall(text)})
    urls = _URL_RE.findall(text)
    linkedin = next((u for u in urls if "linkedin" in u.lower()), None)
    github = next((u for u in urls if "github" in u.lower()), None)
    portfolio = next((u for u in urls
                      if "linkedin" not in u.lower() and "github" not in u.lower()), None)

    low = text.lower()
    skills = sorted({s for s in _SKILL_KEYWORDS
                     if re.search(rf"(?<![\w+#]){re.escape(s)}(?![\w+#])", low)})
    titles = sorted({t for t in _TITLE_KEYWORDS if t in low})

    # Heuristic name: first short, digit-free, '@'-free line near the top.
    name = None
    for line in text.splitlines()[:8]:
        s = line.strip()
        if s and "@" not in s and not re.search(r"\d", s) and 1 < len(s.split()) <= 4:
            name = s
            break

    return {
        "name": name,
        "emails": emails,
        "phones": phones,
        "linkedin_url": linkedin,
        "github_url": github,
        "portfolio_url": portfolio,
        "skills": skills,
        "titles": titles,
    }


# Fields a resume genuinely cannot supply — these must come from the wizard (SPEC §10.13).
_STILL_NEEDED_ALWAYS = [
    "work_authorization", "notice_period", "availability",
    "salary_expectation", "willing_to_relocate", "remote_preference",
]


def cmd_parse_resume(args: argparse.Namespace) -> int:
    src = os.path.abspath(os.path.expanduser(args.path))
    if not os.path.isfile(src):
        print(f"❌ file not found: {src}")
        return 2
    if not src.lower().endswith(".docx"):
        print(f"❌ expected a .docx resume, got: {src}")
        return 2

    hp = home_paths()
    os.makedirs(hp["data"], exist_ok=True)
    dst = os.path.join(hp["data"], os.path.basename(src))
    if os.path.abspath(src) != os.path.abspath(dst):
        shutil.copy2(src, dst)

    try:
        text = read_docx_text(dst)
    except Exception as exc:
        print(f"❌ could not read resume: {exc}")
        return 1

    extracted = extract_from_resume(text)
    still_needed = list(_STILL_NEEDED_ALWAYS)
    if not extracted["phones"]:
        still_needed.append("personal.phone")
    if not extracted["linkedin_url"]:
        still_needed.append("personal.linkedin_url")

    result = {
        "resume_copied_to": dst,
        "extracted": extracted,
        "still_needed": still_needed,
    }
    print(json.dumps(result, indent=2, ensure_ascii=False))
    return 0


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="alfred",
        description="Alfred setup & health CLI. Secrets enter ONLY via `set-secret` (getpass).")
    sub = parser.add_subparsers(dest="command")

    sub.add_parser("bootstrap",
                   help="create the venv + install deps + init (cross-platform, no shell)")
    sub.add_parser("env", help="print the venv's python path for this OS")

    sub.add_parser("init", help="create ~/.alfred/{config,data,output}/ + copy templates")

    p_secret = sub.add_parser("set-secret", help="getpass (hidden) → write KEY into .env")
    p_secret.add_argument("key", help="e.g. TELEGRAM_BOT_TOKEN, GMAIL_ADDRESS, GMAIL_APP_PASSWORD")

    sub.add_parser("detect-telegram-chat-id",
                   help="read getUpdates → write TELEGRAM_CHAT_ID into .env")

    sub.add_parser("doctor", help="full readiness checklist (superset of --preflight)")

    p_resume = sub.add_parser("parse-resume", help="copy resume into data/ + report fields")
    p_resume.add_argument("path", help="path to your master resume .docx")

    return parser


def main(argv: Optional[List[str]] = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    dispatch: Dict[str, Callable[[argparse.Namespace], int]] = {
        "bootstrap": cmd_bootstrap,
        "env": cmd_env,
        "init": cmd_init,
        "set-secret": cmd_set_secret,
        "detect-telegram-chat-id": cmd_detect,
        "doctor": cmd_doctor,
        "parse-resume": cmd_parse_resume,
    }
    handler = dispatch.get(args.command)
    if handler is None:
        parser.print_help()
        return 2
    return handler(args)


if __name__ == "__main__":
    raise SystemExit(main())
