"""WarmApply · Resume + Cover Letter DRY-RUN (truthfulness core demo).

Proves the format-safe tailoring end-to-end, WITHOUT touching data/ or the network:
  1. build a small ATS-simple SAMPLE resume.docx in a TEMP folder (never data/),
  2. take one FAKE enriched job,
  3. extract paragraphs, tailor ONLY the summary + re-emphasize 1-2 EXISTING bullets
     (no invented skills — every keyword used is already in the sample resume),
  4. write the tailored .docx via scripts/resume_edit.py,
  5. attempt PDF via scripts/docx_to_pdf.py (report gracefully if LibreOffice absent),
  6. generate a cover letter + short email from FIXED skeletons,
  7. print a "what I changed" diff PROVING only the targeted paragraphs differ
     (all other paragraphs byte-identical in text AND style).

Everything is labeled [FAKE]/[SAMPLE]. No live calls. Deps: python-docx only.
"""

from __future__ import annotations

import os
import sys
import tempfile

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import docx_to_pdf  # noqa: E402
import resume_edit  # noqa: E402


# ---------------------------------------------------------------------------
# FAKE enriched job (company-research output shape) + FAKE profile.
# ---------------------------------------------------------------------------

FAKE_JOB = {
    "job_id": "linkedin:FAKE-A1",
    "title": "Site Reliability Engineer",
    "company": "Globex Systems",
    "url": "https://example.com/fake/A1",
    "description_snippet": (
        "[FAKE] SRE role: own SLOs and on-call. Stack is Kubernetes on AWS, Docker, "
        "Terraform for IaC, and Python automation across Linux fleets. Go is a plus."
    ),
    "company_analysis": {
        "recent_signal": "announced a multi-region platform launch",
        "values_culture": "reliability-first, strong on-call culture",
        "website": "https://globex.example",
    },
    "email": {"address": "priya.sharma@globex.example", "recipient_name": "Priya Sharma"},
    "match": {"score": 91},
}

FAKE_PROFILE = {
    "personal": {
        "full_name": "Fake Candidate",
        "email": "fake.candidate@example.com",
        "phone": "+91-90000-00000",
        "linkedin_url": "https://linkedin.com/in/fake-candidate",
    },
    "cover_letter_tone": "warm, concise, direct",
}


# ---------------------------------------------------------------------------
# Build a small ATS-simple sample resume (single column, standard styles).
# The skills it re-emphasizes below are ALL present here — nothing is invented.
# ---------------------------------------------------------------------------

def build_sample_resume(path: str) -> None:
    doc = Document()
    doc.add_paragraph("Fake Candidate", style="Title")                       # 0
    doc.add_paragraph("Bengaluru, India · fake.candidate@example.com")       # 1
    doc.add_paragraph("Professional Summary", style="Heading 1")             # 2
    doc.add_paragraph(                                                        # 3  <- SUMMARY (target)
        "Infrastructure engineer with experience keeping cloud services "
        "running and automating operational work.")
    doc.add_paragraph("Experience", style="Heading 1")                       # 4
    doc.add_paragraph("Cloud Operations Engineer — Initrode (2021–2026)")    # 5
    doc.add_paragraph(                                                        # 6  <- BULLET (target)
        "Managed container orchestration and infrastructure-as-code for "
        "cloud services.", style="List Bullet")
    doc.add_paragraph(                                                        # 7  <- BULLET (target)
        "Wrote automation scripts to reduce manual toil and speed up "
        "deployments.", style="List Bullet")
    doc.add_paragraph(                                                        # 8  (untouched bullet)
        "Handled on-call rotations and incident response for production "
        "systems.", style="List Bullet")
    doc.add_paragraph("Skills", style="Heading 1")                           # 9
    doc.add_paragraph(                                                        # 10 (real skills)
        "Python, AWS, Docker, Kubernetes, Terraform, Linux, CI/CD")
    doc.add_paragraph("Education", style="Heading 1")                        # 11
    doc.add_paragraph("B.E. Computer Science — Fake University (2021)")      # 12
    doc.save(path)


# ---------------------------------------------------------------------------
# Fixed skeletons (truthful fill only).
# ---------------------------------------------------------------------------

def cover_letter_text(job, profile) -> str:
    p = profile["personal"]
    name = p["full_name"]
    recipient = job["email"].get("recipient_name") or "Hiring Team"
    company = job["company"]
    role = job["title"]
    hook = job["company_analysis"].get("recent_signal", "your work")
    return (
        f"Dear {recipient},\n\n"
        f"I'm excited to apply for the {role} role at {company}. Your recent "
        f"{hook} is exactly the kind of reliability challenge I enjoy.\n\n"
        f"In my work I've managed container orchestration and infrastructure-as-code "
        f"for cloud services and written automation to cut operational toil — "
        f"experience that maps directly to owning SLOs and on-call for your platform.\n\n"
        f"I'd welcome the chance to discuss how I can help keep {company}'s systems "
        f"reliable. Thank you for your consideration.\n\n"
        f"Warm regards,\n{name}\n{p['email']} · {p['phone']}"
    )


def email_text(job, profile) -> str:
    p = profile["personal"]
    first = (job["email"].get("recipient_name") or "there").split()[0]
    role = job["title"]
    company = job["company"]
    subject = f"{role} — {p['full_name']}"
    body = (
        f"Subject: {subject}\n\n"
        f"Hi {first},\n\n"
        f"I saw the {role} opening at {company} — your recent "
        f"{job['company_analysis'].get('recent_signal', 'work')} caught my eye. "
        f"I've spent the last few years managing Kubernetes orchestration and "
        f"Terraform-based infrastructure for cloud services, which lines up well "
        f"with owning reliability here.\n\n"
        f"I've attached my resume and a short cover letter.\n\n"
        f"Best,\n{p['full_name']}\n{p['email']} · {p['linkedin_url']}"
    )
    return body


# ---------------------------------------------------------------------------
# Dry-run
# ---------------------------------------------------------------------------

def main() -> int:
    print("=" * 74)
    print("WarmApply · Resume + Cover Letter — DRY RUN (truthfulness core)")
    print("=" * 74)

    workdir = tempfile.mkdtemp(prefix="warmapply_resume_")
    outdir = os.path.join(workdir, "output",
                          f"{FAKE_JOB['company'].replace(' ', '')}_"
                          f"{FAKE_JOB['title'].replace(' ', '')}")
    os.makedirs(outdir, exist_ok=True)

    master = os.path.join(workdir, "sample_master_resume.docx")
    build_sample_resume(master)
    print(f"\n[SAMPLE] master resume : {master}")
    print(f"[FAKE]   job            : {FAKE_JOB['title']} @ {FAKE_JOB['company']}")
    print(f"output dir             : {outdir}")

    # --- 1. Extract paragraphs (the agent reads this to pick targets) ---
    before = resume_edit.extract(master)
    print("\nExtracted paragraphs (index · style · text):")
    for row in before:
        txt = row["text"] if len(row["text"]) <= 60 else row["text"][:57] + "..."
        print(f"  [{row['index']:>2}] {str(row['style']):<12} {txt}")

    # --- 2. Decide truthful edits (summary + 2 existing bullets) ---
    # Every keyword below is already present in the sample resume's Skills line
    # (Python, AWS, Docker, Kubernetes, Terraform, Linux, CI/CD) — nothing invented.
    edits = {
        3: ("Site Reliability Engineer with hands-on experience running "
            "Kubernetes on AWS and automating infrastructure with Terraform and "
            "Python to keep cloud services reliable."),
        6: ("Managed Kubernetes container orchestration and Terraform "
            "infrastructure-as-code across AWS cloud services."),
        7: ("Wrote Python automation and CI/CD scripts to reduce manual toil "
            "and speed up deployments."),
    }
    tailored_docx = os.path.join(outdir, f"Resume_{FAKE_JOB['company'].replace(' ', '')}.docx")
    resume_edit.apply_edits(master, edits, tailored_docx)
    print(f"\nTailored .docx written : {tailored_docx}")

    # --- 3. Attempt PDF (graceful if LibreOffice missing) ---
    print("\nPDF conversion:")
    if docx_to_pdf.is_available():
        try:
            pdf = docx_to_pdf.convert(tailored_docx, outdir)
            print(f"  ✅ Resume PDF: {pdf}")
        except Exception as e:
            print(f"  ⚠️  conversion failed: {e}")
    else:
        print("  ⚠️  LibreOffice (soffice) not installed — skipping PDF.")
        print("     Install: macOS `brew install --cask libreoffice` / "
              "Ubuntu `sudo apt-get install libreoffice`")
        print("     (.docx is still produced; PDF is generated when soffice exists.)")

    # --- 4. Cover letter + email (fixed skeletons) ---
    cl_docx = os.path.join(outdir, f"CoverLetter_{FAKE_JOB['company'].replace(' ', '')}.docx")
    cl_doc = Document()
    for para in cover_letter_text(FAKE_JOB, FAKE_PROFILE).split("\n"):
        cl_doc.add_paragraph(para)
    cl_doc.save(cl_docx)
    email_path = os.path.join(outdir, "email_message.txt")
    with open(email_path, "w", encoding="utf-8") as fh:
        fh.write(email_text(FAKE_JOB, FAKE_PROFILE))
    print(f"\nCover letter .docx     : {cl_docx}")
    print(f"Email message          : {email_path}")

    # --- 5. Verify format lock: only targeted paragraphs changed ---
    after = resume_edit.extract(tailored_docx)
    targeted = set(edits.keys())
    print("\n" + "-" * 74)
    print("FORMAT-LOCK PROOF — per-paragraph comparison (master vs tailored):")
    all_ok = True
    changed_report = []
    for b, a in zip(before, after):
        idx = b["index"]
        text_changed = b["text"] != a["text"]
        style_same = b["style"] == a["style"]
        if idx in targeted:
            # Expected to change text, but MUST keep its style.
            status = "CHANGED (style preserved)" if style_same else "CHANGED — STYLE BROKEN!"
            if not (text_changed and style_same):
                all_ok = False
            changed_report.append((idx, b["text"], a["text"], style_same))
        else:
            # Must be byte-identical: same text AND same style.
            if text_changed or not style_same:
                status = "UNEXPECTED CHANGE!"
                all_ok = False
            else:
                status = "identical"
        flag = "•" if idx in targeted else " "
        print(f"  {flag}[{idx:>2}] {status}")

    print("\nBefore → after (only the 3 targeted paragraphs):")
    for idx, old, new, style_ok in changed_report:
        print(f"\n  [{idx}] style preserved: {style_ok}")
        print(f"      BEFORE: {old}")
        print(f"      AFTER : {new}")

    # --- 6. what_i_changed.md (mandatory, must match the diff) ---
    wic = os.path.join(outdir, "what_i_changed.md")
    lines = [
        f"# What I changed — {FAKE_JOB['title']} @ {FAKE_JOB['company']}",
        "",
        "Only the professional summary and two existing bullets were re-worded. "
        "Format, projects, titles, dates, and every other paragraph are unchanged.",
        "",
        "## Edits (before → after)",
    ]
    labels = {3: "Summary rewrite", 6: "Bullet re-emphasis", 7: "Bullet re-emphasis"}
    for idx, old, new, _ in changed_report:
        lines += [f"### [{idx}] {labels.get(idx, 'Edit')}",
                  f"- **Before:** {old}", f"- **After:** {new}", ""]
    lines += [
        "## Keywords surfaced (all already in the master resume)",
        "- Kubernetes, AWS, Terraform, Python, CI/CD — present in the Skills line; "
        "re-emphasized, not invented.",
        "",
        "## Gaps NOT filled (truthfulness)",
        "- The JD lists **Go** as a plus. It is NOT in the master resume, so it was "
        "deliberately left out rather than fabricated.",
    ]
    with open(wic, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines) + "\n")
    print(f"\nwhat_i_changed.md      : {wic}")

    print("\n" + "=" * 74)
    verdict = "PASS ✅ — only targeted paragraphs changed; all others identical." if all_ok \
        else "FAIL ❌ — unexpected changes detected!"
    print(f"Truthfulness / format-lock check: {verdict}")
    print("=" * 74)
    return 0 if all_ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
