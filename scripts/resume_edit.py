"""WarmApply · resume_edit — format-safe DOCX paragraph editing (python-docx only).

This is the mechanical half of WarmApply's truthfulness core. The agent (the brain)
decides WHICH paragraphs are the summary / which bullets to re-emphasize and supplies
the new truthful text; this module performs the edit WITHOUT touching anything else.

Guarantees:
  - Only the paragraph at the given index changes.
  - That paragraph keeps its paragraph STYLE and its first run's formatting
    (font, size, bold/italic, color) — text is swapped inside the existing run.
  - Every other paragraph (and all images/tables) is left as-is.

Public API (per the resume-cover-letter definition):
    - extract(docx_path) -> [{index, style, text}]
    - replace_paragraph_text(docx_path, index, new_text, out_path) -> out_path

Convenience:
    - apply_edits(docx_path, edits, out_path) -> out_path   # {index: new_text}

Dependency: python-docx (already in requirements.txt). No network.

Note: operates on body-level paragraphs (`document.paragraphs`), which is the
right scope for ATS-simple, single-column resumes (no layout tables/text boxes).
"""

from __future__ import annotations

from typing import Dict, List

from docx import Document


def extract(docx_path: str) -> List[Dict[str, object]]:
    """Return every body paragraph as {index, style, text}, in document order."""
    doc = Document(docx_path)
    out: List[Dict[str, object]] = []
    for i, para in enumerate(doc.paragraphs):
        out.append({
            "index": i,
            "style": para.style.name if para.style is not None else None,
            "text": para.text,
        })
    return out


def _set_paragraph_text_preserving_format(para, new_text: str) -> None:
    """Replace a paragraph's text, keeping paragraph style + first run's format.

    Strategy: reuse the first run (so its <w:rPr> formatting carries over), set its
    text to `new_text`, and drop any remaining runs so no stale text lingers. The
    paragraph's own style/properties (<w:pPr>) are never touched.
    """
    runs = para.runs
    if runs:
        first = runs[0]
        first.text = new_text
        # Remove the other runs' XML elements (their text is now redundant).
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        # Empty paragraph (e.g. a spacer) — add a run; nothing to preserve.
        para.add_run(new_text)


def replace_paragraph_text(docx_path: str, index: int, new_text: str,
                           out_path: str) -> str:
    """Replace the text of ONE paragraph (by index), preserving its formatting.

    Reads `docx_path`, edits paragraph `index`, writes to `out_path`, and returns
    `out_path`. All other paragraphs, images, and tables are left intact.
    """
    doc = Document(docx_path)
    paras = doc.paragraphs
    if index < 0 or index >= len(paras):
        raise IndexError(
            f"paragraph index {index} out of range (0..{len(paras) - 1})")
    _set_paragraph_text_preserving_format(paras[index], new_text)
    doc.save(out_path)
    return out_path


def apply_edits(docx_path: str, edits: Dict[int, str], out_path: str) -> str:
    """Apply several {index: new_text} edits in a single load/save pass.

    Equivalent to calling replace_paragraph_text once per index, but avoids
    intermediate files. Only the named indices change.
    """
    doc = Document(docx_path)
    paras = doc.paragraphs
    n = len(paras)
    for index, new_text in edits.items():
        if index < 0 or index >= n:
            raise IndexError(
                f"paragraph index {index} out of range (0..{n - 1})")
        _set_paragraph_text_preserving_format(paras[index], new_text)
    doc.save(out_path)
    return out_path
